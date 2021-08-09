import xlrd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# 统计功能点数量
count = 0


map_function = set()

def add_func_point(word_doc, style=None,heading_str="无", define_str="无", requirement_str="无", element_in_str="无", element_out_str="无"):
    global count
    count = count + 1

    if heading_str in map_function:
        return

    map_function.add(heading_str)

    word_doc.add_heading("",6).add_run(heading_str,  style=style)
    template = """
【业务定义】
    {}	
【功能要求】
    {}
【业务要素】
    输入要素：{}
    输出要素：{}
    """.format(define_str, requirement_str, element_in_str, element_out_str)
    word_doc.add_paragraph("").add_run(template, style=style)


def add_sheet(sheet, word, paragraph_style, heading_style):
    # 功能需求
    cur_col_c = ""
    # 触发事件列
    cur_col_d = ""
    # 功能过程列
    cur_col_e = ""
    # 子过程描述列
    cur_col_f = ""
    # 数据组
    cur_col_h = ""
    # 数据属性
    cur_col_i = ""

    for i in range(sheet.nrows):
        if i == 0:
            continue
        row_values = sheet.row_values(i)
        if len(row_values) < 8:
            continue

        if row_values[0]:
            word.add_heading("", 3).add_run(row_values[0], style=heading_style)
        if row_values[2]:
            cur_col_c = row_values[2]
            word.add_heading("", 4).add_run(cur_col_c, style=heading_style)
        if row_values[3]:
            cur_col_d = row_values[3]
        if row_values[4]:
            cur_col_e = row_values[4]
            # word.add_heading("", 5).add_run(cur_col_e, style=paragraph_style)
        if row_values[5]:
            cur_col_f = row_values[5]
        if row_values[7]:
            cur_col_h = row_values[7]
        if row_values[8]:
            cur_col_i = row_values[8]
        print("功能需求:{}, 触发事件:{}, 功能过程:{}, 子过程描述:{}, 数据组:{}, 数据属性:{}".format(
            cur_col_c, cur_col_d, cur_col_e, cur_col_f, cur_col_h, cur_col_i
        ))

        heading_str = cur_col_f
        define_str = "{},{}".format(cur_col_c, cur_col_e)
        requirement_str = "{},{},{}".format(cur_col_c, cur_col_e, cur_col_f)
        element_in_str = "{},{}".format(cur_col_h, cur_col_i)
        element_out_str = "{},{}".format(cur_col_h, cur_col_i)
        add_func_point(word, style=paragraph_style,heading_str=heading_str, define_str=define_str, requirement_str=requirement_str,
                       element_in_str=element_in_str, element_out_str=element_out_str)


def get_font_style(word, name, font_name, font_size):
    # 创建自定义字符样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    user_font_style = word.styles.add_style(name, 2)
    # 设置字体尺寸
    user_font_style.font.size = Pt(font_size)
    # 设置字体颜色
    user_font_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # 设置段落样式为宋体
    user_font_style.font.name = font_name
    user_font_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return user_font_style


if __name__ == '__main__':
    word = Document()
    paragraph_style = get_font_style(word, 'UserStyle2', '宋体', 10.5)
    heading_style = get_font_style(word, 'UserStyle3', '黑体', 14)

    wb = xlrd.open_workbook(filename="/Users/yejiaxin/Desktop/cosmic/中移动内蒙公司计费帐务中心COSMIC20210115_new.xlsx")
    # sheet = wb.sheet_by_name("cosmic平台-VOS")

    for sheet_name in ['cosmic平台-VOS', 'cosmic平台 DSM', 'cosmic平台 DSE', 'cosmic平台 XC']:
    # for sheet_name in ['cosmic平台-VOS']:
        sheet = wb.sheet_by_name(sheet_name)
        add_sheet(sheet, word, paragraph_style, heading_style)

    word.save("/Users/yejiaxin/Desktop/cosmic/云南移动计费账务中心V8项目-需求规格书.docx")

    print("总共{}个功能点".format(count))
